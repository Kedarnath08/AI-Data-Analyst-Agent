"""PDF/DOCX ingestion tests. Fixtures are generated at test time so no binary
sample files need to live in the repo."""
import docx
import pytest

from src.ingest.docx_tables import ingest_docx
from src.ingest.pdf_tables import ingest_pdf


def _make_docx_with_table(path):
    d = docx.Document()
    d.add_paragraph("Quarterly report")
    table = d.add_table(rows=3, cols=2)
    data = [("Quarter", "Sales"), ("Q1", "100"), ("Q2", "200")]
    for r, (a, b) in enumerate(data):
        table.rows[r].cells[0].text = a
        table.rows[r].cells[1].text = b
    d.save(str(path))


def _make_docx_text_only(path):
    d = docx.Document()
    d.add_paragraph("First paragraph.")
    d.add_paragraph("Second paragraph.")
    d.save(str(path))


def test_docx_with_table(tmp_path):
    p = tmp_path / "report.docx"
    _make_docx_with_table(p)
    result = ingest_docx("docx1", p)
    assert result["extraction_method"] == "python_docx_tables"
    assert result["tables"][0]["name"] == "table_1"
    assert result["tables"][0]["row_count"] == 2
    cols = {c["name"] for c in result["tables"][0]["columns"]}
    assert cols == {"quarter", "sales"}


def test_docx_text_fallback(tmp_path):
    p = tmp_path / "notes.docx"
    _make_docx_text_only(p)
    result = ingest_docx("docx2", p)
    assert result["extraction_method"] == "text_fallback"
    assert result["tables"][0]["name"] == "document_text"
    assert result["tables"][0]["row_count"] == 2


# --- PDF: build a tiny PDF with a table using reportlab if available ---
reportlab = pytest.importorskip("reportlab", reason="reportlab not installed; skip PDF gen")


def _make_pdf_with_table(path):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle

    doc = SimpleDocTemplate(str(path), pagesize=letter)
    data = [["Name", "Score"], ["Alice", "90"], ["Bob", "85"]]
    table = Table(data)
    # Grid lines make the table detectable by pdfplumber's line-based extractor.
    table.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black)]))
    doc.build([table])


def test_pdf_with_table(tmp_path):
    p = tmp_path / "scores.pdf"
    _make_pdf_with_table(p)
    result = ingest_pdf("pdf1", p)
    assert result["extraction_method"] == "pdfplumber_tables"
    assert len(result["tables"]) >= 1
    assert result["tables"][0]["row_count"] >= 2
