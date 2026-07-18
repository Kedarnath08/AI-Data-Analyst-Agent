"""DOCX ingestion: extract Word tables with python-docx, fall back to paragraphs.

Each table in the document becomes its own DuckDB table. When the document has
no usable tables, non-empty paragraphs are stored as a document_text table.
"""
from pathlib import Path

import docx
import pandas as pd

from src import duck
from src.ingest import _describe, _dedupe_cols, _safe_col
from src.ingest.tabular import _load_df


def ingest_docx(dataset_id: str, path: Path) -> dict:
    document = docx.Document(str(path))

    con = duck.connect(dataset_id)
    tables = []
    try:
        for i, tbl in enumerate(document.tables):
            rows = [[cell.text for cell in row.cells] for row in tbl.rows]
            if len(rows) < 2:
                continue
            header, *body = rows
            cols = _dedupe_cols([_safe_col(h, j) for j, h in enumerate(header)])
            df = pd.DataFrame(body, columns=cols)
            table = f"table_{i + 1}"
            _load_df(con, table, df)
            tables.append(_describe(con, table))

        if tables:
            method = "python_docx_tables"
        else:
            paras = [p.text for p in document.paragraphs if p.text.strip()]
            df = pd.DataFrame(
                {"paragraph_index": list(range(len(paras))), "text": paras}
            )
            _load_df(con, "document_text", df)
            tables.append(_describe(con, "document_text"))
            method = "text_fallback"
    finally:
        con.close()
    return {"tables": tables, "extraction_method": method}
